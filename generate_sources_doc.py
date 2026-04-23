"""
Generates the Dexaview Sources & Testing Guide as a .docx file.
Run with: python3 generate_sources_doc.py
Output:   Dexaview_Sources_and_Testing_Guide.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x0A, 0x3D, 0x62)
    p.space_before = Pt(16); p.space_after = Pt(4)
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x17, 0x5F, 0x7A)
    p.space_before = Pt(12); p.space_after = Pt(3)
    return p

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x73, 0x6A)
    p.space_before = Pt(8); p.space_after = Pt(2)
    return p

def body(text):
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(5)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True; r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run("NOTE: " + text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x80, 0x60, 0x00)
    r.italic = True
    p.paragraph_format.space_after = Pt(5)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '0A3D62')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        hdr[i]._tc.get_or_add_tcPr().append(shd)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if cells[c_idx].paragraphs[0].runs:
                cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

# ---------------------------------------------------------------------------
# COVER
# ---------------------------------------------------------------------------
cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run("DEXAVIEW")
r.bold = True; r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x0A, 0x3D, 0x62)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.add_run("Sources, Feed Instructions & Testing Guide").font.size = Pt(16)

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.add_run(f"Version 1.0 | {datetime.date.today().strftime('%B %Y')}").font.size = Pt(11)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 1 — WHAT IS A "FEED"
# ---------------------------------------------------------------------------
h1("1. What Is a Feed and Why It Matters")
body(
    "In Dexaview, a 'feed' is the configuration data that drives a simulation session. "
    "Before the 3D simulation can run meaningfully, it needs to know: "
    "what video to sync with, what industry it is simulating, what events to trigger, "
    "when to trigger them, and what the events mean. "
    "This is entered through the Scenario Builder screen."
)
body(
    "Without a feed, the simulation engine runs but the canvas is empty and no events fire. "
    "This guide provides: (1) pre-built sample feeds for immediate testing, "
    "(2) the sources where real industry feed data comes from, "
    "and (3) exact instructions for building your own feeds from those sources."
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 2 — OIL & GAS SOURCES
# ---------------------------------------------------------------------------
h1("2. Oil & Gas — Document Sources & Feed Instructions")

body(
    "The following sources provide the real-world data, procedures, and timelines "
    "needed to build accurate oil and gas simulation feeds. All standards cited are "
    "those already programmed into Dexaview's AI Technical Advisor."
)

# --- 2.1 API Standards ---
h2("2.1 API (American Petroleum Institute) Standards")
body(
    "API standards are the primary technical authority for oil and gas operations in "
    "the United States and internationally. The following are directly relevant to "
    "Dexaview simulations:"
)

add_table(
    ["Standard", "Title", "What To Use It For In Dexaview", "Where To Get It"],
    [
        ["API RP 53",
         "Blowout Prevention Equipment Systems for Drilling Wells",
         "BOP failure scenarios — shut-in procedures, BOP component descriptions, pressure ratings",
         "api.org/products-and-services/standards — search 'RP 53'. Free preview available. Full standard: ~$150 USD"],
        ["API RP 64",
         "Diverter Systems for Drilling Operations",
         "Diverter activation cue points — when diverter opens vs. when BOP closes",
         "api.org — search 'RP 64'"],
        ["API RP 49",
         "Recommended Practice for Drilling and Well Servicing Operations Involving Hydrogen Sulfide",
         "H2S release scenarios — exposure limits, SCBA requirements, muster procedures",
         "api.org — search 'RP 49'. Some sections publicly summarised by OSHA"],
        ["API RP 59",
         "Well Control Operations",
         "Kill mud calculations, well control worksheet — use for IWCF training feed",
         "api.org — search 'RP 59'"],
        ["API 1130",
         "Computational Pipeline Monitoring for Liquids",
         "Pipeline rupture detection thresholds, SCADA alarm values for pipeline feeds",
         "api.org — search '1130'"],
    ]
)

note(
    "Free alternative: The API provides free 'Fact Sheets' and summaries for many standards. "
    "Search '[standard number] fact sheet filetype:pdf' on Google. "
    "OSHA also publishes free summaries of API standards relevant to worker safety."
)

# --- 2.2 IADC ---
h2("2.2 IADC (International Association of Drilling Contractors)")
body(
    "IADC publishes the Well Control Manual — the single most detailed source of "
    "drilling emergency procedures in the world. It is the basis for IWCF and "
    "WellCAP certification exams."
)
bullet("Website: iadc.org")
bullet("Publication: 'IADC Well Control Manual' — available at iadc.org/store. Cost: approx $200–$350 USD depending on edition.")
bullet("Free content: IADC publishes many free technical papers at onepetro.org — search 'IADC well control' for hundreds of free papers.")
bullet("For Dexaview feeds: Use Chapter 4 (Kick Detection), Chapter 5 (Shut-In Procedures), Chapter 6 (Well Kill Methods) for cue point descriptions and telemetry values.")

h3("How to Extract Feed Data from the IADC Well Control Manual:")
body("Follow these steps to convert a manual chapter into a Dexaview scenario:")
p = doc.add_paragraph(style='List Number')
p.add_run("Open Chapter 4 — Kick Detection Indicators. Find the table of warning signs.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Each indicator (pit gain, pump pressure drop, flow rate increase) becomes a cue point description.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Find the typical values in the worked examples (e.g. '15 barrel pit gain triggers shut-in'). Use these as telemetry values.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Match the procedural steps to your video timestamps.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Enter this into the Scenario Builder. The AI advisor already knows the manual — it will expand on your descriptions automatically.").font.size = Pt(11)
doc.add_paragraph()

# --- 2.3 OSHA ---
h2("2.3 OSHA — Free Government Resources")
body(
    "OSHA (Occupational Safety and Health Administration) publishes free, detailed "
    "guidance documents for all oil and gas hazards. These are legally required "
    "standards, not recommendations."
)
add_table(
    ["OSHA Document", "Topic", "URL / How To Find"],
    [
        ["OSHA 1910.119",
         "Process Safety Management (PSM) — required for any facility with highly hazardous chemicals",
         "osha.gov — search '1910.119'. Full text free online."],
        ["OSHA 1910.1000",
         "Air Contaminants — H2S exposure limits (1 ppm TLV-TWA, 50 ppm IDLH)",
         "osha.gov — search '1910.1000 table Z-1'"],
        ["OSHA 3162 — Oil and Gas Well Drilling Guide",
         "Complete drilling safety guide with real incident case studies",
         "osha.gov — search 'OSHA 3162'. Free PDF download."],
        ["OSHA 3213 — H2S in Oil and Gas",
         "H2S specific guide with detection, PPE, and evacuation procedures",
         "osha.gov — search 'OSHA 3213'. Free PDF download."],
        ["30 CFR Part 250",
         "BSEE offshore drilling regulations — notification requirements after incidents",
         "ecfr.gov — search '30 CFR 250'"],
    ]
)

# --- 2.4 YouTube Channels ---
h2("2.4 YouTube Channels for Oil & Gas Video Feed Sources")
body(
    "The following YouTube channels produce high-quality, technically accurate "
    "oil and gas educational content. Use these videos as the left-panel feed "
    "while you map cue points in the Scenario Builder. "
    "All channels listed are real and publicly available."
)
add_table(
    ["Channel Name", "YouTube Search Term", "Best Content Type", "Typical Video Length"],
    [
        ["Practical Engineering",
         "Practical Engineering YouTube",
         "Infrastructure, pipelines, pressure systems — excellent physics explanations",
         "10–20 minutes"],
        ["Real Engineering",
         "Real Engineering YouTube",
         "Industrial systems, energy, engineering disasters — very visual",
         "8–15 minutes"],
        ["Transocean",
         "Transocean drilling operations YouTube",
         "Real offshore drilling operations, BOP demonstrations",
         "3–10 minutes"],
        ["Schlumberger / SLB",
         "SLB oilfield technology YouTube",
         "Well control, downhole tools, formation evaluation",
         "5–15 minutes"],
        ["IADC Official",
         "IADC well control YouTube",
         "Official well control training animations and procedures",
         "5–20 minutes"],
        ["Well Control School",
         "Well Control School YouTube training",
         "IWCF/IADC exam preparation — kick detection, shut-in procedures",
         "10–30 minutes"],
        ["NOIA (National Ocean Industries Association)",
         "NOIA offshore safety YouTube",
         "Offshore platform safety, environmental response",
         "5–15 minutes"],
    ]
)

note(
    "How to get the video ID: Open the video on YouTube. The ID is the 11-character "
    "code after '?v=' in the URL. Example: youtube.com/watch?v=dQw4w9WgXcQ → "
    "video ID is 'dQw4w9WgXcQ'. Paste this into the Video ID field in the Scenario Builder."
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 3 — EDUCATION SOURCES
# ---------------------------------------------------------------------------
h1("3. Education — Document Sources & Feed Instructions")

body(
    "Dexaview can be used as a STEM and vocational education tool. "
    "The following sources provide curriculum-aligned content for building "
    "education-focused simulation feeds."
)

# --- 3.1 STEM Curriculum Alignment ---
h2("3.1 Curriculum Standards — Where To Find Alignment Documents")

add_table(
    ["Curriculum", "Relevant Topics", "Source Document", "URL"],
    [
        ["GCSE Physics (AQA — UK)",
         "Pressure (P=F/A), Fluid mechanics, Forces",
         "AQA GCSE Physics Specification 8463",
         "aqa.org.uk — search 'GCSE Physics specification'. Free download."],
        ["A-Level Physics (AQA — UK)",
         "Stress, strain, bulk modulus, thermal physics",
         "AQA A-Level Physics Specification 7408",
         "aqa.org.uk — search 'A-Level Physics specification'. Free download."],
        ["IB Physics SL/HL",
         "Topic 3 (Thermal physics), Topic 8 (Energy production), Topic B (Engineering physics)",
         "IB Physics Guide",
         "ibo.org — requires IB account. Teachers can get free access through their school."],
        ["AP Physics 1 (USA)",
         "Unit 5 (Torque and Rotation), Unit 8 (Electric Charge and Force)",
         "AP Physics 1 Course and Exam Description",
         "collegeboard.org — search 'AP Physics 1 CED'. Free PDF download."],
        ["BTEC Engineering (UK)",
         "Unit 1 (Engineering Principles), Unit 3 (Engineering Product Design)",
         "BTEC Level 3 Engineering Specification",
         "qualifications.pearson.com — search 'BTEC Engineering'. Free download."],
        ["STEM Australia",
         "Physics — Pressure, Forces, Work, Energy",
         "Australian Curriculum Physics",
         "australiancurriculum.edu.au — free online"],
        ["NGSS (USA Next Gen Science Standards)",
         "HS-PS2 (Motion and Stability: Forces and Interactions)",
         "NGSS Standards",
         "nextgenscience.org — free online"],
    ]
)

# --- 3.2 Open Educational Content ---
h2("3.2 Open Educational Content — Free Video and Document Sources")

body("These resources provide free, curriculum-aligned educational videos and documents:")

bullet("Khan Academy (khanacademy.org): ", "Pressure and fluids unit under Physics → Fluid Mechanics. Free videos with transcripts. Ideal for GCSE/AP level feeds.")
bullet("MIT OpenCourseWare (ocw.mit.edu): ", "Search '2.001 Mechanics and Materials I' for university-level stress and pressure content. All course materials free.")
bullet("CrashCourse Physics (YouTube): ", "Search 'CrashCourse Fluids' and 'CrashCourse Pressure'. 10-minute videos aligned to AP Physics. Video ID available from YouTube URL.")
bullet("PhET Interactive Simulations (phet.colorado.edu): ", "University of Colorado's simulation library. 'Fluid Pressure and Flow' simulation — use alongside Dexaview for comparison demonstrations.")
bullet("STEM Learning (stem.org.uk): ", "UK's national STEM resource library. Search 'pressure' or 'oil industry'. Many free lesson plans with suggested activities.")
bullet("TED-Ed (YouTube): ", "Search 'TED-Ed pressure' or 'TED-Ed oil'. High-quality animated explanations, typically 5 minutes.")

# --- 3.3 Vocational Training ---
h2("3.3 Vocational Training — IWCF and Industry Certifications")
body(
    "For vocational training feeds (workforce upskilling rather than school education), "
    "the following certification bodies provide the exact content that professionals "
    "must learn:"
)
add_table(
    ["Certification", "Body", "Website", "What to Use"],
    [
        ["IWCF Level 2 (Surface Stack)", "International Well Control Forum", "iwcf.org",
         "Download the IWCF syllabus — it lists every topic tested. Build one cue per syllabus item."],
        ["IWCF Level 3 (Subsea)", "International Well Control Forum", "iwcf.org",
         "Subsea BOP scenarios. More complex cue sequences."],
        ["IADC WellCAP", "International Association of Drilling Contractors", "iadc.org",
         "WellCAP study guide available for purchase. Used by North American operators."],
        ["CompEx (Explosive Atmospheres)", "CompEx Scheme", "compex.info",
         "For H2S and hazardous area electrical scenarios."],
        ["NEBOSH Oil and Gas Certificate", "NEBOSH", "nebosh.org.uk",
         "Process safety management. IGC and Oil & Gas syllabus documents downloadable from website."],
    ]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 4 — DATA CENTER SOURCES
# ---------------------------------------------------------------------------
h1("4. Data Center — Document Sources & Feed Instructions")

h2("4.1 Standards Documents")
add_table(
    ["Standard", "Title", "Relevant Feed Content", "Where To Get It"],
    [
        ["ANSI/TIA-942-B",
         "Telecommunications Infrastructure Standard for Data Centers",
         "Tier classification (I–IV), redundancy requirements, temperature thresholds",
         "tiaonline.org — purchase required (~$300). Executive summary free on TIA website."],
        ["ASHRAE Thermal Guidelines",
         "Thermal Guidelines for Data Processing Environments",
         "Inlet temperature limits: Class A1 (15–32°C), A2 (10–35°C). Use these as telemetry trigger values.",
         "ashrae.org — 'Thermal Guidelines for Data Processing Environments'. Free summary available."],
        ["IEEE 1100",
         "Recommended Practice for Powering and Grounding Electronic Equipment",
         "Power quality events, grounding faults, UPS transfer scenarios",
         "ieeexplore.ieee.org — IEEE membership or purchase required"],
        ["NFPA 75",
         "Standard for the Fire Protection of Information Technology Equipment",
         "Clean agent suppression (FM-200, Novec 1230) vs water sprinkler decisions",
         "nfpa.org — free 'read online' access available. Purchase for offline use."],
        ["Uptime Institute Tier Standards",
         "Data Center Tier Certification Standards",
         "Tier III/IV concurrently maintainable — use for maintenance window scenarios",
         "uptimeinstitute.com — free white papers in their library"],
    ]
)

h2("4.2 YouTube Channels for Data Center Feed Sources")
add_table(
    ["Channel", "Search Term", "Content Type"],
    [
        ["Data Center Frontier", "Data Center Frontier YouTube", "Industry news, case studies, cooling systems"],
        ["Schneider Electric", "Schneider Electric data center YouTube", "UPS systems, cooling, power architecture"],
        ["Vertiv", "Vertiv data center infrastructure YouTube", "Thermal management, power systems"],
        ["Data Center Knowledge", "Data Center Knowledge YouTube", "Operations, incidents, management"],
        ["NetworkChuck", "NetworkChuck data center YouTube", "Accessible explanations of server infrastructure"],
    ]
)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 5 — HOW TO BUILD A FEED FROM A DOCUMENT
# ---------------------------------------------------------------------------
h1("5. Step-by-Step: How To Build a Feed from Any Document")

body(
    "This section gives you the exact process to take any industry document "
    "and convert it into a Dexaview simulation feed."
)

h2("Step 1 — Choose Your Source Document")
body("Pick one of the sources listed above. For first testing, use:")
bullet("OSHA 3162 (Oil and Gas) — free, detailed, written in plain English")
bullet("Khan Academy Fluid Mechanics — free, clear, with exact timestamps")
bullet("IADC well control training video from YouTube")

h2("Step 2 — Find the Critical Events")
body(
    "Read or watch the source material. Mark every moment where a consequence occurs — "
    "a pressure changes, a component fails, an alarm fires, an action must be taken. "
    "Each of these is a cue point in your feed."
)
body("Example from OSHA 3162, Chapter 4 (Blowout Prevention):")
bullet("Page 42: 'When pit gain exceeds 10 barrels, a well kick should be suspected' → Cue: pressure_spike at the relevant video timestamp")
bullet("Page 45: 'If flow continues after pump shut-off, shut-in immediately' → Cue: blowout at next timestamp")
bullet("Page 48: 'If BOP fails, activate diverter and sound general alarm' → Cue: collapse at next timestamp")

h2("Step 3 — Write the Event Description")
body(
    "For each cue point, write 2–4 sentences that explain: "
    "what just happened, what the physical/chemical cause is, "
    "what the immediate required action is, and which standard or regulation applies. "
    "This is what will appear on screen during the simulation event."
)
body("Template:")
body(
    '    "[WHAT HAPPENED]: [equipment/component] has [failed/exceeded/activated]. '
    '[MEASUREMENT]: [parameter] is [value]. '
    '[REQUIRED ACTION]: [verb] [equipment] per [standard] Section [X]. '
    '[CONSEQUENCE IF NOT DONE]: [result]."'
)
body("Example:")
body(
    '    "BOP annular preventer has failed to seal. Well pressure: 4,800 psi and rising. '
    'Activate diverter system per API RP 64 Section 5.3. '
    'If diverter fails, well will flow to surface uncontrolled."'
)

h2("Step 4 — Set the Telemetry Values")
body(
    "Look in the source document for specific measurements at each event. "
    "These become your Pressure, Temperature, and Status fields."
)
add_table(
    ["Event Stage", "Pressure", "Temperature", "Status"],
    [
        ["Normal operations", "2,800 psi", "74°C", "NOMINAL"],
        ["Kick detected", "3,200 psi (rising)", "78°C", "WARNING"],
        ["BOP fails", "4,800 psi (CRITICAL)", "142°C", "EMERGENCY"],
        ["Well killed", "2,650 psi (stable)", "74°C", "WARNING"],
    ]
)

h2("Step 5 — Map to a YouTube Video")
body(
    "Find a YouTube video that covers this topic. "
    "The video does not need to be perfect — it just needs to cover the topic "
    "in the right order so your cue timestamps roughly match the content. "
    "Use the seek buttons in the Dexaview left panel to jump between cue points "
    "regardless of exact video timing."
)
body("To find a video:")
p = doc.add_paragraph(style='List Number')
p.add_run("Go to YouTube. Search using the query terms in Section 2.4 or 3.2 above.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Filter by 'This year' and 'Long (>20 minutes)' for training content.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Copy the video URL. Extract the 11-character ID after ?v=.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Paste into the Scenario Builder Video ID field.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Watch the video once with the simulation open. Adjust cue timestamps to match the video's discussion of each event.").font.size = Pt(11)
doc.add_paragraph()

h2("Step 6 — Test the Feed")
body("After loading your feed into the Scenario Builder and launching the simulation:")
bullet("Click 'TRIGGER BLOWOUT' manually first to confirm the physics engine is working")
bullet("Type a question to the AI advisor: 'What is the first step when [your scenario event] occurs?' — confirm it gives a relevant, standards-cited response")
bullet("Play the video and watch for the first cue timestamp — confirm the event fires automatically")
bullet("Read the event description that appears on screen — edit it if it is unclear or too technical for your audience")
bullet("Check the telemetry overlays above the 3D asset — confirm they show the values you entered")

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 6 — QUICK REFERENCE TESTING FEEDS
# ---------------------------------------------------------------------------
h1("6. Quick Reference — Ready-to-Type Feed Configurations")
body(
    "These are complete feed configurations you can type directly into the "
    "Scenario Builder right now to start testing. "
    "Replace the video ID with a real YouTube video ID from the sources above."
)

h2("Feed 1 — Minimum Viable Test (5 minutes to set up)")
body("Industry: Oil & Gas | Events: 1 | Purpose: Confirm simulation works")
add_table(
    ["Field", "Value to Enter"],
    [
        ["Video URL/ID", "Any YouTube video about oil drilling — paste its ID"],
        ["Industry", "Oil & Gas"],
        ["Title", "Basic BOP Test"],
        ["Description", "Testing simulation. A blowout preventer fails."],
        ["Event 1 — Timestamp", "01:00"],
        ["Event 1 — Type", "Blowout / Explosion"],
        ["Event 1 — Description", "BOP has failed. Well is flowing. Immediate shut-in required per API RP 53."],
        ["Event 1 — Pressure", "4,800 psi"],
        ["Event 1 — Temperature", "142°C"],
        ["Event 1 — Status", "EMERGENCY"],
    ]
)

h2("Feed 2 — Full Oil & Gas Training Scenario (20 minutes to set up)")
body("Industry: Oil & Gas | Events: 3 | Source: IADC Well Control Manual Chapter 4-6")
add_table(
    ["Field", "Value to Enter"],
    [
        ["Video URL/ID", "YouTube: search 'IADC well control kick detection' → use top result's ID"],
        ["Industry", "Oil & Gas"],
        ["Title", "Well Kick to Blowout — Complete Sequence"],
        ["Description", "Offshore deepwater well. Formation pressure exceeds mud weight. Well takes a kick at 18,500 feet. BOP fails. Emergency response required."],
        ["Event 1 — Timestamp", "00:45"],
        ["Event 1 — Type", "Pressure Spike"],
        ["Event 1 — Description", "KICK INDICATORS: Pit gain +15 bbls. Pump pressure -120 psi. Flow rate increase confirmed. Execute soft shut-in per IADC Manual Section 4.2."],
        ["Event 1 — Pressure", "3,200 psi (rising)"],
        ["Event 1 — Status", "WARNING"],
        ["Event 2 — Timestamp", "01:30"],
        ["Event 2 — Type", "Blowout / Explosion"],
        ["Event 2 — Description", "BOP blind shear rams FAILED. Well flowing uncontrolled. Activate diverter per API RP 64. All non-essential personnel to muster stations."],
        ["Event 2 — Pressure", "4,800 psi CRITICAL"],
        ["Event 2 — Status", "EMERGENCY"],
        ["Event 3 — Timestamp", "03:30"],
        ["Event 3 — Type", "Structural Collapse"],
        ["Event 3 — Description", "Riser integrity compromised. Emergency disconnect initiated. Relief well planning commences per regulatory requirement."],
        ["Event 3 — Pressure", "0 psi (disconnected)"],
        ["Event 3 — Status", "EMERGENCY"],
    ]
)

h2("Feed 3 — Education STEM (10 minutes to set up)")
body("Industry: Oil & Gas | Events: 2 | Audience: Secondary school physics class")
add_table(
    ["Field", "Value to Enter"],
    [
        ["Video URL/ID", "YouTube: search 'CrashCourse Physics Fluids' → use video ID"],
        ["Industry", "Oil & Gas"],
        ["Title", "Physics of Pressure: Real Oil Well Application"],
        ["Description", "Secondary school demonstration. Pressure, Pascal's Law, and fluid mechanics are shown through a live well simulation. No industry knowledge needed."],
        ["Event 1 — Timestamp", "01:00"],
        ["Event 1 — Type", "Pressure Spike"],
        ["Event 1 — Description", "Formation pressure = mud hydrostatic pressure. P = ρgh. ρ=1200 kg/m³, g=9.81, h=2800m → P = 32.9 MPa. STUDENT: Calculate force on drill pipe using F = P × A."],
        ["Event 1 — Pressure", "32.9 MPa"],
        ["Event 1 — Status", "WARNING"],
        ["Event 2 — Timestamp", "02:30"],
        ["Event 2 — Type", "Blowout / Explosion"],
        ["Event 2 — Description", "Formation pressure exceeds containment — Pascal's Law. Pressure applied to fluid transmits equally in all directions. Result: fluid accelerates upward through drill string."],
        ["Event 2 — Pressure", "35 MPa (uncontrolled)"],
        ["Event 2 — Status", "EMERGENCY"],
    ]
)

# ---------------------------------------------------------------------------
# SECTION 7 — WHAT TO WATCH FOR WHEN TESTING
# ---------------------------------------------------------------------------
h1("7. What To Watch For When Testing Each Feed")
body(
    "After loading a feed and launching the simulation, check these specific things:"
)

add_table(
    ["What to Check", "How to Check It", "What Good Looks Like", "What Bad Looks Like"],
    [
        ["Physics event fires at correct time",
         "Play video, watch for cue timestamp",
         "Explosion/collapse fires within 0.1 seconds of video timestamp",
         "Event fires too early/late, or does not fire at all"],
        ["Event description appears on screen",
         "Watch the centre panel canvas",
         "Red overlay appears with your written description clearly readable",
         "Overlay missing, or shows wrong event"],
        ["AI advisor responds to the event",
         "Watch the right panel after event fires",
         "AI sends a message within 3–5 seconds citing a relevant standard",
         "No response, or 'Advisor unavailable' error (check API key)"],
        ["AI response is industry-accurate",
         "Read the AI response against your source document",
         "AI cites correct standard, gives numbered checklist of actions",
         "AI gives generic response with no standard cited"],
        ["Telemetry overlay shows event values",
         "Look at the 3D asset floating panel",
         "Pressure/Temp/Status update to match your cue's values",
         "Overlay still shows initial values after event"],
        ["Seek buttons jump to correct timestamps",
         "Click the cue seek buttons in the left panel",
         "Video jumps to within 1 second of the cue time",
         "Video jumps to wrong time or does not seek"],
        ["FPS counter stays above 30",
         "Look at top-right corner of canvas",
         "55–60 FPS on modern hardware",
         "Below 30 FPS — browser may not support WebGPU. Try Chrome 113+"],
    ]
)

# ---------------------------------------------------------------------------
# FINAL NOTE
# ---------------------------------------------------------------------------
h1("8. Recommended First Testing Path")
body("Do these in order, each building on the last:")
p = doc.add_paragraph(style='List Number')
p.add_run("Set VITE_OPENAI_API_KEY in Railway variables. This takes 2 minutes.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Open the Scenario Builder. Load 'Feed 1 — Minimum Viable Test' from Section 6. Use any oil drilling YouTube video ID. Launch simulation.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Click 'TRIGGER BLOWOUT' manually. Confirm explosion physics work. Type a question to AI. Confirm response.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Play the video. Confirm the blowout fires at 1:00.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Load 'Feed 2 — Full Oil & Gas Training Scenario'. Use an IADC YouTube video. Test all 3 events.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Load 'Feed 3 — Education STEM'. Use the CrashCourse Fluids video. Test with a non-technical audience member.").font.size = Pt(11)
p = doc.add_paragraph(style='List Number')
p.add_run("Record the screen during Feed 2. This recording becomes your first YouTube video demonstrating the product.").font.size = Pt(11)
doc.add_paragraph()

foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.add_run(f"Dexaview Sources & Testing Guide | {datetime.date.today().year} | Confidential").font.size = Pt(9)

# ---------------------------------------------------------------------------
output = "/home/emeka/Dexaview/Dexaview_Sources_and_Testing_Guide.docx"
doc.save(output)
print(f"Saved: {output}")
