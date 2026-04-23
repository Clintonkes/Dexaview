"""
Generates the Dexaview partner briefing document as a .docx file.
Run with:  python3 generate_doc.py
Output:    Dexaview_Partner_Briefing.docx  (in the same directory)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ---------------------------------------------------------------------------
# Styles helpers
# ---------------------------------------------------------------------------

def set_heading(paragraph, text, level=1):
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x0A, 0x3D, 0x62)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x17, 0x5F, 0x7A)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    paragraph.space_before = Pt(14)
    paragraph.space_after = Pt(4)

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text).font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '175F7A')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:val'), 'clear')
        hdr[i].paragraphs[0]._p.get_or_add_pPr().append(shading)
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

# ---------------------------------------------------------------------------
# COVER PAGE
# ---------------------------------------------------------------------------

doc.add_picture  # placeholder — skip image, use text cover

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("DEXAVIEW")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0x0A, 0x3D, 0x62)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("The Industrial Metaverse for YouTube")
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x17, 0x5F, 0x7A)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub2.add_run("Complete Partner & Co-Founder Briefing Document")
r3.font.size = Pt(13)
r3.italic = True

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}").font.size = Pt(11)

doc.add_page_break()

# ---------------------------------------------------------------------------
# SECTION 1 — EXECUTIVE SUMMARY
# ---------------------------------------------------------------------------

h = doc.add_paragraph()
set_heading(h, "1. Executive Summary", 1)

body(
    "Dexaview is a browser-based industrial simulation platform that transforms "
    "YouTube educational videos into live, interactive 3D simulations. It is the "
    "first platform to synchronise a YouTube video timestamp with a real-time "
    "physics simulation engine — a technology called the Sim-Link."
)
body(
    "When a viewer watches a Dexaview video on YouTube and clicks the Sim-Link "
    "in the description, their browser opens a three-panel interface: the YouTube "
    "video on the left, a live 3D industrial simulation in the centre (an oil rig, "
    "a data centre, a pipeline), and an AI Technical Advisor on the right powered "
    "by GPT-4o trained on industry safety standards."
)
body(
    "As the video plays, the simulation reacts in real time. When the instructor "
    "describes a blowout preventer failure at 1 minute 30 seconds, the 3D rig "
    "explodes with accurate physics at that exact moment. The AI immediately "
    "explains the correct emergency procedure, citing API RP 53 and the IADC Well "
    "Control Manual."
)
body("The business operates across four revenue streams rated Very High feasibility:")
bullet("B2B Training Licensing — energy firms pay $500–$5,000/month for private branded instances")
bullet("YouTube Ad Revenue — industrial content earns $8–$25 CPM vs. $2–$4 for general content")
bullet("Premium Sim-Link Subscriptions — $9.99/month for saved reports, full scenario access")
bullet("Creator Marketplace — 30% commission on 3D asset and simulation logic sales")

spacer()

# ---------------------------------------------------------------------------
# SECTION 2 — HOW THE APPLICATION WORKS (TECHNICAL)
# ---------------------------------------------------------------------------

h = doc.add_paragraph()
set_heading(h, "2. How The Application Works — Technical Architecture", 1)

h2 = doc.add_paragraph()
set_heading(h2, "2.1 The Three-Panel Interface", 2)

body(
    "Every user sees three panels side by side when they open the Sim-Link URL:"
)
bullet("LEFT PANEL — YouTube Player: The video plays here. Quick-seek buttons jump to simulation cue points.")
bullet("CENTRE PANEL — 3D Simulation Canvas: A live WebGPU-rendered industrial environment with real physics (Rapier physics engine at 60Hz). Floating telemetry overlays show Pressure, Temperature, and Status above each asset in real time.")
bullet("RIGHT PANEL — AI Technical Advisor: A chat interface powered by GPT-4o. The AI is pre-loaded with industry safety standards and automatically responds when events trigger in the simulation.")

h2 = doc.add_paragraph()
set_heading(h2, "2.2 The Sim-Link — How Video and Simulation Connect", 2)

body(
    "The Sim-Link is the core innovation. It works through a system called cue points:"
)
numbered("The creator registers timestamps against event names. Example: at 1 minute 30 seconds, trigger a 'blowout' event.")
numbered("A polling engine reads the YouTube player's current time 60 times per second.")
numbered("When the video reaches 1:30, the simulation fires the blowout event automatically.")
numbered("The 3D physics engine applies a blast-radius impulse — objects within 15 metres of the origin fly outward with force that decreases with distance.")
numbered("Simultaneously, the AI advisor is asked: 'A blowout event has just been triggered. What is the immediate response procedure?' It answers with a numbered checklist referencing API RP 53.")
numbered("The viewer can also click on 3D objects to seek the video backward — if they click the BOP stack, the video jumps to the moment it was discussed.")

h2 = doc.add_paragraph()
set_heading(h2, "2.3 The Scenario Builder — User-Driven Simulation Input", 2)

body(
    "Before a simulation runs, users need to be able to describe what they want "
    "to simulate. This is the Scenario Builder — a pre-simulation configuration "
    "screen that collects:"
)
bullet("YouTube Video URL or ID — the video the simulation will sync with")
bullet("Industry Mode — Oil & Gas, Data Center, Pipeline, Offshore, Manufacturing")
bullet("Scenario Title and Description — written in plain language by the user")
bullet("Event Cue Points — each cue has: a timestamp, an event type (blowout, collapse, fire, flood), and a plain-English description shown to viewers during the event")
bullet("Custom Telemetry Values — the pressure, temperature, and status readings shown on the 3D asset overlays")

body(
    "For example, a trainer at an oil company would fill in: "
    "Video URL: [their internal training video], Industry: Oil & Gas, "
    "Scenario: 'Deepwater Blowout During Hurricane Season', "
    "Cue at 2:15: blowout — description: 'Riser pressure exceeds 10,000 PSI. BOP fails to seal. "
    "Immediate well control procedure required.' The simulation then runs exactly as configured."
)

h2 = doc.add_paragraph()
set_heading(h2, "2.4 The Physics Engine — What 'Real' Simulation Means", 2)

body(
    "The simulation uses Rapier — a production-grade physics engine written in "
    "Rust and compiled to WebAssembly, running at 60 physics steps per second in "
    "the browser with no download required. This means:"
)
bullet("Blowout event: A spherical blast force is calculated. Every dynamic object within 15 metres receives an outward impulse scaled by distance from the origin. Closer objects fly faster.")
bullet("Collapse event: Objects above the failure origin receive a downward impulse of 500 Newton-seconds, simulating structural failure.")
bullet("Trimesh colliders: The 3D model geometry is used directly as the physics collision surface — not simplified boxes. The rig physically behaves like the rig looks.")
bullet("All simulations run at Earth gravity (9.81 m/s²) by default. Underwater simulations can reduce this to simulate buoyancy.")

h2 = doc.add_paragraph()
set_heading(h2, "2.5 The AI Technical Advisor — Industry Knowledge Built In", 2)

body(
    "The AI advisor is GPT-4o configured with deep industry knowledge loaded at "
    "startup. It knows the following standards and always cites them in responses:"
)

body("For Oil & Gas simulations:")
bullet("API RP 53 — Blowout Prevention Equipment Systems")
bullet("API RP 64 — Diverter Systems for Drilling Operations")
bullet("IADC Well Control Manual")
bullet("OSHA 1910.119 — Process Safety Management of Highly Hazardous Chemicals")

body("For Data Center simulations:")
bullet("ANSI/TIA-942 — Data Centre Infrastructure Standard")
bullet("ASHRAE A2 thermal guidelines")
bullet("IEEE 1100 — Powering and Grounding Sensitive Electronic Equipment")
bullet("NFPA 75 — Fire Protection of Information Technology Equipment")

body(
    "The advisor maintains a rolling conversation history of up to 20 turns so "
    "follow-up questions are answered in context. It can be switched to data center "
    "mode, pipeline mode, or any other industry by changing a single parameter."
)

spacer()

# ---------------------------------------------------------------------------
# SECTION 3 — SIMULATION REPORTS
# ---------------------------------------------------------------------------

h = doc.add_paragraph()
set_heading(h, "3. Simulation Reports — What Gets Generated", 1)

body(
    "After every simulation session, a structured report is generated. This is "
    "the primary deliverable for B2B training customers. A report includes:"
)

add_table(
    ["Report Field", "What It Contains", "Industry Use"],
    [
        ["Session ID & Timestamp", "Unique ID, date, time of session", "Audit trail for compliance"],
        ["Industry Mode", "Oil & Gas / Data Center / Pipeline", "Routes to correct standards"],
        ["Scenario Title", "User-defined name of the scenario", "Training record identification"],
        ["Duration", "Total simulation time in seconds", "Engagement measurement"],
        ["Events Triggered", "Name, timestamp, origin of each physics event", "Procedure compliance check"],
        ["AI Conversation Log", "Full Q&A between user and advisor", "Knowledge assessment"],
        ["Response Time", "Seconds from event trigger to user's first question", "Competency measurement"],
        ["Standards Referenced", "Which API/OSHA/ANSI codes the AI cited", "Compliance documentation"],
        ["Performance", "Average frames per second during session", "Technical quality log"],
    ]
)

body(
    "For an oil company running a blowout prevention training session, the report "
    "shows: did the trainee ask the right questions? Did they ask within the "
    "required response window? Which procedures did they follow? This report is "
    "signed off and stored against the employee's training record — exactly what "
    "HSE regulators require."
)

spacer()

# ---------------------------------------------------------------------------
# SECTION 4 — HOW TO TEST THE APPLICATION
# ---------------------------------------------------------------------------

h = doc.add_paragraph()
set_heading(h, "4. How To Test The Application Right Now", 1)

body("There are three things that must be configured before the application is fully functional:")

h2 = doc.add_paragraph()
set_heading(h2, "Step 1 — Set The OpenAI API Key", 2)
body("In Railway's dashboard: Frontend service → Variables → Add:")
body("    VITE_OPENAI_API_KEY = sk-your-openai-key-here")
body("Without this, the AI advisor panel shows 'Advisor unavailable.'")

h2 = doc.add_paragraph()
set_heading(h2, "Step 2 — Replace The YouTube Video ID", 2)
body("In the source file src/components/SimLinkPage.jsx, line 36, change:")
body("    const DEMO_VIDEO_ID = 'REPLACE_WITH_YOUR_YOUTUBE_VIDEO_ID';")
body("to the ID from your YouTube video URL (the characters after ?v=).")

h2 = doc.add_paragraph()
set_heading(h2, "Step 3 — Add The 3D Industrial Asset", 2)
body("Place a .glb 3D model file at: public/assets/drilling_rig.glb")
body("Free sources: Sketchfab (search 'drilling rig' or 'oil platform'), TurboSquid, or Poly Haven.")
body("This file is automatically served by nginx in production and loaded into the simulation canvas.")

h2 = doc.add_paragraph()
set_heading(h2, "Testing Checklist After Setup", 2)

add_table(
    ["Test", "Action", "Expected Result"],
    [
        ["Physics — Blowout", "Click 'TRIGGER BLOWOUT' button", "All objects within 15m fly outward with blast falloff"],
        ["Physics — Collapse", "Click 'TRIGGER COLLAPSE' button", "Objects above origin drop with 500 Ns downward impulse"],
        ["AI Advisor", "Type any question about well control", "GPT-4o responds citing API RP 53 within 3 seconds"],
        ["Sim-Link Sync", "Play video, wait for 1m30s timestamp", "Blowout triggers automatically, AI responds unprompted"],
        ["Telemetry Overlay", "Load the .glb file", "Floating panel shows Pressure / Temp / Status above asset"],
        ["Reverse Seek", "Click a 3D object", "YouTube video seeks to the relevant timestamp"],
        ["FPS Counter", "Open the app", "FPS displayed on canvas — should be 55–60fps"],
    ]
)

spacer()

# ---------------------------------------------------------------------------
# SECTION 5 — YOUTUBE CHANNEL SETUP
# ---------------------------------------------------------------------------

h = doc.add_paragraph()
set_heading(h, "5. YouTube Channel Setup & Content Strategy", 1)

h2 = doc.add_paragraph()
set_heading(h2, "5.1 Channel Configuration", 2)

body("Channel name: Dexaview  or  Dexaview Live")
body("Channel description:")
body(
    '    "The Industrial Metaverse for YouTube. Watch real industrial crises '
    'unfold — then click the Sim-Link to try it yourself in a live 3D simulation. '
    'No download. No hardware. Just your browser."'
)
bullet("Channel banner: Dark industrial visual (oil rig, server hall, pipeline at night) with the three-panel interface as the hero image")
bullet("Channel URL: Link directly to dexaview-production.up.railway.app")
bullet("Channel trailer: 60-second cut of the most dramatic simulation moment, ending with the CTA")

h2 = doc.add_paragraph()
set_heading(h2, "5.2 The Video Template — Every Video Follows This Structure", 2)

add_table(
    ["Timestamp", "Segment", "Content"],
    [
        ["0:00 – 0:45", "Hook", "Open with the most dramatic moment — rig explosion, server rack fire, pipeline rupture. No intro, no title card. Maximum impact in the first 5 seconds."],
        ["0:45 – 3:00", "Context", "What is this industrial scenario? What are the real-world stakes? Reference an actual incident (Deepwater Horizon, 2021 Texas power grid failure)."],
        ["3:00 – 8:00", "Expert Walkthrough", "What went wrong. What the correct procedure is. Show the equipment involved. This is the training content."],
        ["8:00 – 9:30", "Sim-Link Reveal", "Show the Dexaview interface reacting to this video in real time. The explosion fires in the 3D sim. The AI advisor responds."],
        ["9:30 – 10:00", "CTA", "'Click the Sim-Link in the description. Can you handle this scenario better than our expert? Post your score.'"],
    ]
)

h2 = doc.add_paragraph()
set_heading(h2, "5.3 The Description Template — Copy This For Every Video", 2)

body("Use this exact format in every YouTube video description:")
body("")
body("⚡ DEXAVIEW SIM-LINK: https://dexaview-production.up.railway.app")
body("")
body("Watch what happens at [TIMESTAMP] — then click the Sim-Link to try the same scenario yourself.")
body("The simulation runs live in your browser. No download needed.")
body("")
body("📊 POST YOUR SCORE in the comments.")
body("🏆 Top score this week wins [PRIZE].")
body("")
body("Standards referenced in this episode:")
body("• API RP 53 — Blowout Prevention Equipment")
body("• IADC Well Control Manual")
body("• OSHA 1910.119 — Process Safety Management")

h2 = doc.add_paragraph()
set_heading(h2, "5.4 First 5 Videos To Post — In This Order", 2)

add_table(
    ["#", "Title", "Scenario", "Physics Event", "Why This Order"],
    [
        ["1", "What Really Happens When a BOP Fails", "Deepwater blowout preventer failure", "Blowout at 1m30s", "Highest search volume keyword. Establishes authority."],
        ["2", "Data Center Thermal Runaway — Could You Stop It?", "Server hall overheating cascade", "Collapse at 3m00s", "Reaches tech audience. Proves platform is multi-industry."],
        ["3", "The $150M Mistake: Deepwater Drilling Pressure Error", "Well control miscalculation", "Blowout at 2m15s", "Financial angle drives shares on LinkedIn."],
        ["4", "Pipeline Rupture Response — Do You Know the Protocol?", "Gas pipeline failure and emergency shutdown", "Collapse at 4m00s", "Utility/energy companies share internally for training."],
        ["5", "LIVE: F1 Engineer vs Oil Rig Supervisor — Who Handles Crisis Better?", "Competitive simulation head-to-head", "Both events", "Viral format. Post 60s Shorts clip of explosion for discovery."],
    ]
)

h2 = doc.add_paragraph()
set_heading(h2, "5.5 Posting Schedule", 2)

body("Post one video per week for the first 8 weeks. Do not post daily — industrial content needs time for search indexing and B2B sharing cycles.")
bullet("Tuesday 10am — Upload main video (10 minutes)")
bullet("Wednesday — Upload 60-second YouTube Short (the explosion/crisis moment only)")
bullet("Thursday — Post LinkedIn article: 'The real procedure behind [this week's scenario]'")
bullet("Friday — Post score challenge comment response, engage with top commenters")

spacer()

# ---------------------------------------------------------------------------
# SECTION 6 — REVENUE MODEL
# ---------------------------------------------------------------------------

h = doc.add_paragraph()
set_heading(h, "6. Revenue Model — How Money Is Earned", 1)

add_table(
    ["Stream", "Who Pays", "Price Point", "How To Activate"],
    [
        ["B2B Training Licensing", "Energy companies, data center operators, universities", "$500–$5,000/month", "Demo to HSE training managers. Offer a 2-week free trial of private branded instance."],
        ["YouTube Ad Revenue", "Google (via AdSense)", "$8–$25 CPM (industrial content premium)", "Monetise the channel once 1,000 subscribers and 4,000 watch hours are reached."],
        ["Premium Sim-Link", "Individual users", "$9.99/month", "Add Stripe payment. Gate: saved reports, full scenario access, leaderboard ranking."],
        ["Creator Marketplace", "3D artists, simulation developers", "30% commission on sales", "Open marketplace once 10+ creators are building. Start with curated invite-only."],
    ]
)

body(
    "The fastest path to first revenue is B2B licensing. One phone call to an HSE "
    "training manager at a mid-size oil company, with a live demo, is enough to "
    "close a $1,000/month contract. The demo takes 10 minutes: play the video, "
    "watch the simulation react, ask the AI a well control question, show the "
    "generated report."
)

spacer()

# ---------------------------------------------------------------------------
# SECTION 7 — WHAT STILL NEEDS TO BE BUILT
# ---------------------------------------------------------------------------

h = doc.add_paragraph()
set_heading(h, "7. Build Roadmap — What Is Done vs. What Remains", 1)

h2 = doc.add_paragraph()
set_heading(h2, "7.1 What Is Already Built and Deployed", 2)

bullet("WebGPU 3D rendering engine (Three.js) — runs in browser, no download")
bullet("Rapier physics engine (WASM) — 60Hz physics, blowout and collapse events")
bullet("Sim-Link bridge — YouTube timestamp → physics event synchronisation")
bullet("Reverse Sim-Link — 3D object click → video seek")
bullet("AI Technical Advisor — GPT-4o with oil & gas and data center industry context")
bullet("Telemetry overlays — floating Pressure / Temperature / Status panels on 3D assets")
bullet("Three-panel UI — YouTube / Simulation / AI Advisor layout")
bullet("nginx deployment on Railway — production-ready hosting")
bullet("Draco WASM decoder — compressed 3D model loading")

h2 = doc.add_paragraph()
set_heading(h2, "7.2 What Needs To Be Built Next", 2)

add_table(
    ["Feature", "Where It Goes", "Complexity", "Priority"],
    [
        ["Scenario Builder UI", "New page before simulation loads", "Medium", "CRITICAL — needed for user-driven simulations"],
        ["Event Description Panel", "Centre panel overlay during events", "Low", "HIGH — makes events understandable to viewers"],
        ["Session Report Generator", "DexaviewEngine.generateReport()", "Low", "HIGH — primary B2B deliverable"],
        ["PDF Export", "Backend + frontend download button", "Medium", "HIGH — needed for training records"],
        ["3D Drilling Rig Asset", "public/assets/drilling_rig.glb", "Asset acquisition", "CRITICAL — canvas is empty without it"],
        ["Stripe Subscription", "Backend + SimLinkPage premium gate", "Medium", "HIGH — activates Stream 3 revenue"],
        ["User Accounts / Leaderboard", "Backend + new React page", "Medium", "Medium — needed for gamification"],
        ["More Physics Events", "DexaviewEngine (fire, flood, pressure spike)", "Low", "Medium — expands scenario library"],
        ["Backend API Rebuild", "api/ folder (FastAPI)", "Medium", "HIGH — lost in rearrangement, needed for reports/auth"],
        ["industryMode Switcher", "SimLinkPage dropdown", "Low", "Medium — enables data center and pipeline modes"],
    ]
)

spacer()

# ---------------------------------------------------------------------------
# SECTION 8 — SCENARIO BUILDER (HOW USERS DESCRIBE SIMULATIONS)
# ---------------------------------------------------------------------------

h = doc.add_paragraph()
set_heading(h, "8. The Scenario Builder — User-Driven Simulation Design", 1)

body(
    "Currently the simulation is hardcoded with a demo video and two fixed events. "
    "For the platform to work at scale — where any creator or trainer can build "
    "their own simulation — a Scenario Builder screen must be added before the "
    "simulation loads."
)

h2 = doc.add_paragraph()
set_heading(h2, "8.1 What The Scenario Builder Collects", 2)

add_table(
    ["Field", "Input Type", "Example Value"],
    [
        ["YouTube Video URL", "Text input", "https://youtube.com/watch?v=abc123"],
        ["Industry Mode", "Dropdown", "Oil & Gas / Data Center / Pipeline / Offshore / Manufacturing"],
        ["Scenario Title", "Text input", "Deepwater BOP Failure — Hurricane Conditions"],
        ["Scenario Description", "Paragraph text", "A Category 4 hurricane has caused a riser disconnect. Blowout preventer has failed to seal. Well is flowing. Emergency shut-in required."],
        ["Cue Point — Timestamp", "Time input (mm:ss)", "01:30"],
        ["Cue Point — Event Type", "Dropdown", "blowout / collapse / fire / flood / pressure_spike"],
        ["Cue Point — Description", "Text input", "BOP rams fail to close. Well flow detected at surface. Immediate well kill required."],
        ["Telemetry — Pressure", "Number input", "8,500 psi (rising)"],
        ["Telemetry — Temperature", "Number input", "142°C"],
        ["Telemetry — Status", "Dropdown", "NOMINAL / WARNING / CRITICAL / EMERGENCY"],
    ]
)

h2 = doc.add_paragraph()
set_heading(h2, "8.2 How Event Descriptions Appear During Simulation", 2)

body(
    "When a physics event fires, an overlay panel appears in the centre of the "
    "simulation canvas showing:"
)
bullet("Event name in large red text: 'BLOWOUT TRIGGERED'")
bullet("The description written by the scenario creator: 'BOP rams fail to close. Well flow detected at surface. Immediate well kill required.'")
bullet("A timer counting response time (seconds since event fired)")
bullet("A prompt: 'Ask the AI advisor for the correct procedure →'")

body(
    "This turns a raw physics explosion into a comprehensible training event. "
    "The viewer understands what happened, why it matters, and what they should do."
)

h2 = doc.add_paragraph()
set_heading(h2, "8.3 How To Add This To The Codebase", 2)

body(
    "Add a new React page (ScenarioBuilder.jsx) that renders before SimLinkPage. "
    "On form submission, the configuration is passed as URL parameters or stored "
    "in localStorage and read by SimLinkPage on load. The DEMO_VIDEO_ID and "
    "DEMO_CUES constants in SimLinkPage.jsx are replaced with values read from "
    "this configuration. No backend required for the basic version."
)

spacer()

# ---------------------------------------------------------------------------
# SECTION 9 — SUMMARY
# ---------------------------------------------------------------------------

h = doc.add_paragraph()
set_heading(h, "9. Summary — The Three Things That Matter Most Right Now", 1)

numbered(
    "Get the 3D asset in place. Download a drilling rig .glb from Sketchfab and "
    "place it at public/assets/drilling_rig.glb. Rebuild and push to Railway. "
    "The centre panel will show the simulation. This is the visual that sells the product."
)
numbered(
    "Add the Scenario Builder page. This unlocks user-driven simulations and "
    "makes the platform usable by anyone, not just the developer. Without it, "
    "every new scenario requires a code change."
)
numbered(
    "Record Video 1 and post it. 'What Really Happens When a BOP Fails.' "
    "Record the simulation reacting to your narration. Post it with the Sim-Link "
    "in the description. One video is enough to start the first B2B conversation."
)

body(
    "Everything else — reports, subscriptions, leaderboard, marketplace — is "
    "built on top of these three foundations. The engine is ready. The hosting is "
    "live. The business starts when the content starts."
)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run(
    f"Dexaview — Confidential Partner Briefing | {datetime.date.today().year}"
).font.size = Pt(9)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

output_path = "/home/emeka/Dexaview/Dexaview_Partner_Briefing.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
